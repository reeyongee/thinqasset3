#!/usr/bin/env python3
"""Word letterhead: header + footer as matching high-DPI image bands (inline).

Watermark is a separate behind-text image from the real logo mark.
Builds ivory and all-white paper variants.
"""

from __future__ import annotations

from pathlib import Path

import fitz
import numpy as np
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parent
REPO = ROOT.parent
ASSETS = ROOT / "word-assets"
LOGO_MASK = REPO / "public/thinqasset-assets/thinqasset-logo-reversed.png"
# 1600 dpi A4 bands — sharp enough for print; Word embeds without recompress.
DPI = 1600
GOLD = (182, 160, 130)
SYMBOL_END_X = 316
MASK_W = 1360

VARIANTS = (
    {
        "key": "ivory",
        "paper_css": "#fbfaf7",
        "paper_word": "FBFAF7",
        "assets_subdir": "ivory",
        "sample": ROOT / "ThinqAsset-Letterhead.docx",
        "blank": ROOT / "ThinqAsset-Letterhead-Blank.docx",
    },
    {
        "key": "white",
        "paper_css": "#ffffff",
        "paper_word": "FFFFFF",
        "assets_subdir": "white",
        "sample": ROOT / "ThinqAsset-Letterhead-White.docx",
        "blank": ROOT / "ThinqAsset-Letterhead-White-Blank.docx",
    },
)


def mm_px(mm: float) -> int:
    return int(round(mm * DPI / 25.4))


def px_mm(px: int) -> float:
    return px * 25.4 / DPI


def rasterize_page(paper_css: str, out_dir: Path) -> Image.Image:
    """Print letterhead.html to PDF via Playwright, then high-DPI PNG."""
    out_dir.mkdir(parents=True, exist_ok=True)
    html_path = ROOT / "letterhead.html"
    pdf_path = out_dir / "letterhead-blank.pdf"

    try:
        from playwright.sync_api import sync_playwright

        with sync_playwright() as p:
            browser = p.chromium.launch()
            # Higher DPR improves Chromium's PDF glyph rasterization before MuPDF.
            page = browser.new_page(device_scale_factor=4)
            page.goto(html_path.resolve().as_uri(), wait_until="networkidle")
            page.add_style_tag(
                content=(
                    f":root {{ --paper: {paper_css} !important; }}\n"
                    f".sheet {{ background: {paper_css} !important; }}\n"
                )
            )
            page.evaluate("() => document.fonts.ready")
            page.emulate_media(media="print")
            page.pdf(
                path=str(pdf_path),
                width="210mm",
                height="297mm",
                print_background=True,
                margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
                prefer_css_page_size=True,
            )
            browser.close()
    except Exception as exc:
        print(f"Playwright PDF failed ({exc}); using existing PDF if present")
        if not pdf_path.exists():
            raise

    pdf = fitz.open(pdf_path)
    # alpha=False, annots ignored — clean RGB page for band crops
    pix = pdf[0].get_pixmap(matrix=fitz.Matrix(DPI / 72, DPI / 72), alpha=False)
    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
    img.save(out_dir / "page-background.png", compress_level=1)
    return img


def is_paper(r: int, g: int, b: int) -> bool:
    """Ivory or pure white page (not navy footer / gold rule)."""
    return r > 245 and g > 245 and b > 240


def slice_bands(img: Image.Image, out_dir: Path) -> tuple[float, float]:
    w, h = img.size
    header_h = mm_px(40)
    header = img.crop((0, 0, w, header_h))
    header.save(out_dir / "header.png", compress_level=1)

    # Sustained paper run so gold-rule AA doesn't false-trigger.
    footer_top = h
    run = 0
    first_paper = None
    for y in range(h - 1, 0, -1):
        r, g, b = img.getpixel((w // 2, y))
        if is_paper(r, g, b):
            if first_paper is None:
                first_paper = y
            run += 1
            if run >= max(8, DPI // 75):
                footer_top = first_paper + 1
                break
        else:
            run = 0
            first_paper = None

    footer = img.crop((0, footer_top, w, h))
    footer.save(out_dir / "footer.png", compress_level=1)

    header_mm = px_mm(header_h)
    footer_mm = px_mm(footer.size[1])
    print(f"  header {header_mm:.1f}mm · footer {footer_mm:.1f}mm · {header.size[0]}×{header.size[1]} @ {DPI}dpi")
    return header_mm, footer_mm


def build_watermark(out_dir: Path) -> tuple[float, float]:
    """Real mark, ~58mm, NEAREST upscale — hard edges, brand-accurate."""
    mask = Image.open(LOGO_MASK).convert("RGBA")
    mark_w = int(mask.width * SYMBOL_END_X / MASK_W)
    mark = mask.crop((0, 0, mark_w, mask.height))
    tw = mm_px(58)
    th = int(round(mark.height * (tw / mark.width)))
    mark = mark.resize((tw, th), Image.Resampling.NEAREST)
    arr = np.array(mark)
    a = arr[:, :, 3].astype(np.float32)
    out = np.zeros((arr.shape[0], arr.shape[1], 4), dtype=np.uint8)
    out[:, :, 0] = GOLD[0]
    out[:, :, 1] = GOLD[1]
    out[:, :, 2] = GOLD[2]
    out[:, :, 3] = np.clip(a * 0.72, 0, 195).astype(np.uint8)
    img = Image.fromarray(out, "RGBA")
    img.save(out_dir / "watermark.png", compress_level=1)
    print(f"  watermark {img.size[0]}x{img.size[1]} (~58mm)")
    return px_mm(img.size[0]), px_mm(img.size[1])


# Straight gold rail: below header chrome, above footer (matches HTML .rail)
RAIL_LEFT_MM = 12.0
RAIL_TOP_MM = 48.0
RAIL_BOTTOM_MM = 42.0
RAIL_WIDTH_MM = 0.35


def build_rail(out_dir: Path) -> tuple[float, float, float, float]:
    """Opaque gold vertical rule PNG for Word float (no spark / kink)."""
    height_mm = 297.0 - RAIL_TOP_MM - RAIL_BOTTOM_MM
    # Extra width so Word doesn't crush a sub-pixel line
    tw = max(mm_px(RAIL_WIDTH_MM), 4)
    th = mm_px(height_mm)
    img = Image.new("RGBA", (tw, th), (0, 0, 0, 0))
    # Center a solid gold column
    line_w = max(1, mm_px(RAIL_WIDTH_MM))
    x0 = (tw - line_w) // 2
    for x in range(x0, x0 + line_w):
        for y in range(th):
            img.putpixel((x, y), (*GOLD, 255))
    img.save(out_dir / "rail.png", compress_level=1)
    print(f"  rail {tw}x{th} ({RAIL_WIDTH_MM}mm × {height_mm:.1f}mm)")
    return RAIL_LEFT_MM, RAIL_TOP_MM, RAIL_WIDTH_MM, height_mm


def set_run_font(run, name="Calibri", size=11, color=(30, 37, 45)) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)


def clear_runs(paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag.endswith("}r"):
            paragraph._element.remove(child)


def add_bleed_band(paragraph, path: Path, left_mm: float, right_mm: float) -> None:
    clear_runs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Mm(-left_mm)
    paragraph.paragraph_format.right_indent = Mm(-right_mm)
    paragraph.add_run().add_picture(str(path), width=Mm(210))


def add_watermark(header, assets_dir: Path, width_mm: float, height_mm: float) -> None:
    paragraph = header.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(assets_dir / "watermark.png"), width=Mm(width_mm), height=Mm(height_mm))

    drawing = next(c for c in run._element if c.tag.endswith("}drawing"))
    inline = next(c for c in drawing if c.tag.endswith("}inline"))
    extent = next(c for c in inline if c.tag.endswith("}extent"))
    graphic = next(c for c in inline if c.tag.endswith("}graphic"))
    cx, cy = extent.get("cx"), extent.get("cy")

    ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    ns_pic = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    left_emu = int(Mm(210 - width_mm - 8))
    top_emu = int(Mm(168))

    anchor = parse_xml(
        f'<wp:anchor xmlns:wp="{ns_wp}" xmlns:a="{ns_a}" xmlns:r="{ns_r}" xmlns:pic="{ns_pic}" '
        f'distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240" '
        f'behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{left_emu}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{top_emu}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="900" name="Watermark"/>'
        f'<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        f"</wp:anchor>"
    )
    anchor.append(graphic)
    drawing.replace(inline, anchor)


def add_rail(
    header,
    assets_dir: Path,
    left_mm: float,
    top_mm: float,
    width_mm: float,
    height_mm: float,
) -> None:
    """Straight gold vertical rule, behind text — same geometry as HTML .rail."""
    paragraph = header.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    # Display width slightly wider than hairline so it survives Word zoom
    display_w = max(width_mm, 0.45)
    run.add_picture(str(assets_dir / "rail.png"), width=Mm(display_w), height=Mm(height_mm))

    drawing = next(c for c in run._element if c.tag.endswith("}drawing"))
    inline = next(c for c in drawing if c.tag.endswith("}inline"))
    extent = next(c for c in inline if c.tag.endswith("}extent"))
    graphic = next(c for c in inline if c.tag.endswith("}graphic"))
    cx, cy = extent.get("cx"), extent.get("cy")

    ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    ns_pic = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    left_emu = int(Mm(left_mm))
    top_emu = int(Mm(top_mm))

    anchor = parse_xml(
        f'<wp:anchor xmlns:wp="{ns_wp}" xmlns:a="{ns_a}" xmlns:r="{ns_r}" xmlns:pic="{ns_pic}" '
        f'distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658241" '
        f'behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{left_emu}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{top_emu}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="901" name="GoldRail"/>'
        f'<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        f"</wp:anchor>"
    )
    anchor.append(graphic)
    drawing.replace(inline, anchor)


def build_doc(
    assets_dir: Path,
    header_mm: float,
    footer_mm: float,
    wm_w: float,
    wm_h: float,
    rail_geom: tuple[float, float, float, float],
    *,
    paper_word: str,
    with_sample: bool,
) -> Document:
    left_mm, right_mm = 24.0, 20.0
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Mm(header_mm)
    section.bottom_margin = Mm(footer_mm)
    section.left_margin = Mm(left_mm)
    section.right_margin = Mm(right_mm)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)

    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), paper_word)
    doc.element.insert(0, bg)
    settings = doc.settings.element
    if settings.find(qn("w:displayBackgroundShape")) is None:
        settings.append(OxmlElement("w:displayBackgroundShape"))
    if settings.find(qn("w:doNotAutoCompressPictures")) is None:
        settings.append(OxmlElement("w:doNotAutoCompressPictures"))

    header = section.header
    header.is_linked_to_previous = False
    add_bleed_band(header.paragraphs[0], assets_dir / "header.png", left_mm, right_mm)
    add_watermark(header, assets_dir, wm_w, wm_h)
    add_rail(header, assets_dir, *rail_geom)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    add_bleed_band(fp, assets_dir / "footer.png", left_mm, right_mm)

    if with_sample:
        p = doc.add_paragraph()
        set_run_font(p.add_run("Dear Sir or Madam,"))
        for text in (
            "Thank you for your correspondence. This letter confirms our intention to proceed "
            "with the mandate discussed in our recent meeting, subject to the terms set out in "
            "the accompanying documentation.",
            "We look forward to continuing our collaboration and remain at your disposal for "
            "any clarification you may require.",
        ):
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(10)
            para.paragraph_format.space_before = Pt(6)
            set_run_font(para.add_run(text))
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        set_run_font(para.add_run("Yours faithfully,"))
        for line in ("", "", "", "[Name]", "[Title]", "ThinqAsset Fund Management Ltd"):
            para = doc.add_paragraph()
            if line:
                set_run_font(para.add_run(line))
    else:
        doc.add_paragraph()

    return doc


def build_variant(variant: dict) -> None:
    out_dir = ASSETS / variant["assets_subdir"]
    print(f"\n[{variant['key']}] paper {variant['paper_css']}")
    img = rasterize_page(variant["paper_css"], out_dir)
    header_mm, footer_mm = slice_bands(img, out_dir)
    wm_w, wm_h = build_watermark(out_dir)
    rail_geom = build_rail(out_dir)

    # Also mirror primary ivory bands into ASSETS root for quick previews
    if variant["key"] == "ivory":
        for name in ("header.png", "footer.png", "watermark.png", "rail.png", "page-background.png"):
            src = out_dir / name
            if src.exists():
                (ASSETS / name).write_bytes(src.read_bytes())

    sample = variant["sample"]
    build_doc(
        out_dir, header_mm, footer_mm, wm_w, wm_h, rail_geom,
        paper_word=variant["paper_word"], with_sample=True,
    ).save(sample)
    print(f"  Wrote {sample.name} ({sample.stat().st_size // 1024}KB)")

    blank = variant["blank"]
    build_doc(
        out_dir, header_mm, footer_mm, wm_w, wm_h, rail_geom,
        paper_word=variant["paper_word"], with_sample=False,
    ).save(blank)
    print(f"  Wrote {blank.name} ({blank.stat().st_size // 1024}KB)")


def main() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    for variant in VARIANTS:
        build_variant(variant)


if __name__ == "__main__":
    main()
